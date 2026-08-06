import io
import re
from pathlib import Path
from typing import Protocol
from urllib.parse import quote

from docx import Document

_UNSAFE_ASCII_RE = re.compile(r'[\x00-\x1f\x7f"\\]')


class TimedText(Protocol):
    start_ms: int
    end_ms: int
    text: str


def ms_to_srt_timestamp(ms: int) -> str:
    hours = ms // 3_600_000
    minutes = (ms % 3_600_000) // 60_000
    seconds = (ms % 60_000) // 1000
    millis = ms % 1000
    return f"{hours:02d}:{minutes:02d}:{seconds:02d},{millis:03d}"


def build_txt(segments: list[TimedText]) -> str:
    return "\n".join(s.text for s in segments)


def build_srt(segments: list[TimedText]) -> str:
    blocks = [
        f"{idx}\n{ms_to_srt_timestamp(s.start_ms)} --> {ms_to_srt_timestamp(s.end_ms)}\n{s.text}"
        for idx, s in enumerate(segments, start=1)
    ]
    return "\n\n".join(blocks)


def _stem(filename: str) -> str:
    return Path(filename).stem or filename


def build_docx(original_filename: str, segments: list[TimedText]) -> bytes:
    doc = Document()
    doc.add_heading(_stem(original_filename), level=1)
    for s in segments:
        doc.add_paragraph(s.text)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_filename(original_filename: str, ext: str) -> str:
    return f"{_stem(original_filename)}.{ext}"


def content_disposition(filename: str) -> str:
    """Build a Content-Disposition header safe against a hostile original_filename.

    filename is user-controlled (comes straight from the upload's original
    filename), so both the ASCII fallback and the RFC 5987 UTF-8 value must
    be sanitized against quote-breaking and header-injection characters.
    """
    ascii_only = filename.encode("ascii", "ignore").decode("ascii")
    ascii_fallback = _UNSAFE_ASCII_RE.sub("_", ascii_only) or "export"
    utf8_value = quote(filename, safe="")
    return f'attachment; filename="{ascii_fallback}"; filename*=UTF-8\'\'{utf8_value}'
